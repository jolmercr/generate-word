#!/usr/bin/env python3
"""
Script para personalizar el documento de transferencia tecnológica
"""

from docx import Document
import glob
import os
from datetime import datetime

# ============================================
# CONFIGURACIÓN - EDITAR ESTOS VALORES
# ============================================

SISTEMA_NOMBRE = "MiSistemaWeb"
SISTEMA_VERSION = "v1.0.0"
RESPONSABLE = "Juan Pérez"
SERVIDOR_IP = "192.168.1.100"
SERVIDOR_USUARIO = "deployuser"
RUTA_PROYECTO = "/opt/misistema"
DOMINIO = "misistema.com"
BD_PRINCIPAL = "bd_principal"
SERVIDOR_HOSTNAME = "servidor-prod-01"

# ============================================
# MAPEO DE REEMPLAZOS
# ============================================

REEMPLAZOS = {
    '[NOMBRE DEL SISTEMA WEB]': SISTEMA_NOMBRE,
    '[vX.X.X]': SISTEMA_VERSION,
    '[Tu Nombre]': RESPONSABLE,
    '[Tu Nombre Completo]': RESPONSABLE,
    '[IP-servidor]': SERVIDOR_IP,
    f'usuario@[IP-servidor]': f'{SERVIDOR_USUARIO}@{SERVIDOR_IP}',
    '[nombre-servidor]': SERVIDOR_HOSTNAME,
    '/ruta/a/tu/proyecto': RUTA_PROYECTO,
    '[tudominio.com]': DOMINIO,
    '[basedatos]': BD_PRINCIPAL,
    '[Especificar servidores/cloud]': 'Servidor físico dedicado - DC Principal',
    '[Ej: 99.5% en horario comercial]': '99.9% (24/7)',
    '[Ej: RHEL 8.6 / Ubuntu 22.04]': 'Ubuntu 22.04 LTS',
    '[XX GB - ej: 32GB]': '32GB',
    '[XX GB - ej: 500GB]': '500GB SSD',
    '[Especificaciones - ej: 8 cores]': '8 cores @ 3.2GHz',
    '[X.X.X - ej: 4.6.1]': '4.6.1',
    '[XXX.XXX.XXX.XXX]': SERVIDOR_IP,
    '[/opt/aplicacion]': RUTA_PROYECTO,
    '[puerto]': '22',
    '[Especificar acuerdos de nivel de servicio]': '99.9% uptime / Respuesta < 2s',
}

def encontrar_documento():
    """Encuentra el documento de transferencia más reciente"""
    archivos = glob.glob('Transferencia_Tecnologica_*.docx')
    
    if not archivos:
        print("❌ Error: No se encontró ningún documento de transferencia")
        print("   Ejecuta primero: python src/generate_word.py")
        return None
    
    # Obtener el más reciente
    archivo_mas_reciente = max(archivos, key=os.path.getmtime)
    print(f"📄 Documento encontrado: {archivo_mas_reciente}")
    return archivo_mas_reciente

def reemplazar_en_parrafo(paragraph, reemplazos):
    """Reemplaza texto en un párrafo manteniendo el formato"""
    for old_text, new_text in reemplazos.items():
        if old_text in paragraph.text:
            # Guardar el formato del párrafo
            inline = paragraph.runs
            
            # Reemplazar texto
            for run in inline:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

def personalizar_documento(archivo_entrada):
    """Personaliza el documento con los valores configurados"""
    
    print(f"\n🔧 Personalizando documento...")
    
    try:
        # Cargar documento
        doc = Document(archivo_entrada)
        
        reemplazos_realizados = 0
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            texto_original = paragraph.text
            reemplazar_en_parrafo(paragraph, REEMPLAZOS)
            if paragraph.text != texto_original:
                reemplazos_realizados += 1
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        texto_original = paragraph.text
                        reemplazar_en_parrafo(paragraph, REEMPLAZOS)
                        if paragraph.text != texto_original:
                            reemplazos_realizados += 1
        
        # Generar nombre del archivo
        nombre_limpio = SISTEMA_NOMBRE.replace(' ', '_')
        version_limpia = SISTEMA_VERSION.replace('.', '_')
        nuevo_nombre = f'Transferencia_{nombre_limpio}_{version_limpia}_personalizado.docx'
        
        # Guardar documento personalizado
        doc.save(nuevo_nombre)
        
        print(f"\n✅ Personalización completada!")
        print(f"📁 Documento personalizado: {nuevo_nombre}")
        print(f"🔄 Reemplazos realizados: {reemplazos_realizados}")
        print(f"📊 Tamaño: {os.path.getsize(nuevo_nombre) / 1024:.1f} KB")
        
        return nuevo_nombre
        
    except Exception as e:
        print(f"❌ Error al personalizar documento: {e}")
        return None

def mostrar_recordatorios():
    """Muestra recordatorios de qué revisar manualmente"""
    
    print("\n" + "="*60)
    print("📋 RECORDATORIOS - Revisar manualmente:")
    print("="*60)
    
    recordatorios = [
        "1. Insertar diagramas arquitectónicos en la sección 2.1",
        "2. Completar todos los 25 microservicios en la tabla (actualmente solo 5)",
        "3. Verificar/actualizar comandos específicos de despliegue",
        "4. Agregar contactos de emergencia",
        "5. Revisar sección de mejoras planeadas/deuda técnica",
        "6. Actualizar campos de firmas con nombres reales",
        "7. Verificar configuraciones específicas de seguridad",
        "8. Revisar procedimientos de incidentes",
        "9. Actualizar información de backup según infraestructura real",
        "10. Validar métricas de monitoreo"
    ]
    
    for recordatorio in recordatorios:
        print(f"   {recordatorio}")
    
    print("="*60)

def main():
    """Función principal"""
    
    print("🚀 Script de Personalización de Documento de Transferencia")
    print("="*60)
    
    # Encontrar documento
    archivo = encontrar_documento()
    
    if not archivo:
        return 1
    
    # Mostrar configuración
    print("\n🔧 Configuración actual:")
    print(f"   Sistema: {SISTEMA_NOMBRE} ({SISTEMA_VERSION})")
    print(f"   Responsable: {RESPONSABLE}")
    print(f"   Servidor: {SERVIDOR_HOSTNAME} ({SERVIDOR_IP})")
    print(f"   Ruta proyecto: {RUTA_PROYECTO}")
    print(f"   Dominio: {DOMINIO}")
    
    # Personalizar
    nuevo_archivo = personalizar_documento(archivo)
    
    if nuevo_archivo:
        mostrar_recordatorios()
        print("\n✨ ¡Documento listo para revisión final!")
        return 0
    else:
        return 1

if __name__ == "__main__":
    exit(main())
