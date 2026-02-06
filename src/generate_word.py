"""
Generador de Documento de Transferencia Tecnológica
para Sistemas con Microservicios .NET 8 + Podman
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

def crear_documento_transferencia():
    """Crea el documento completo de transferencia tecnológica"""
    
    # Crear documento
    doc = Document()
    
    # Configurar propiedades del documento
    doc.core_properties.title = "Documento de Transferencia Tecnológica"
    doc.core_properties.author = "Departamento de Desarrollo"
    doc.core_properties.subject = "Transferencia Sistema Web Microservicios"
    
    # Configurar estilos personalizados
    configurar_estilos(doc)
    
    # Agregar contenido
    agregar_portada(doc)
    agregar_indice(doc)
    agregar_resumen_ejecutivo(doc)
    agregar_arquitectura(doc)
    agregar_infraestructura(doc)
    agregar_despliegue(doc)
    agregar_configuracion(doc)
    agregar_monitoreo(doc)
    agregar_backup(doc)
    agregar_seguridad(doc)
    agregar_incidentes(doc)
    agregar_mejoras(doc)
    agregar_anexos(doc)
    agregar_firmas(doc)
    
    # Guardar documento
    nombre_archivo = f"Transferencia_Tecnologica_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(nombre_archivo)
    print(f"✅ Documento generado: {nombre_archivo}")
    print(f"📁 Ruta: {os.path.abspath(nombre_archivo)}")
    
    return nombre_archivo

def configurar_estilos(doc):
    """Configura los estilos personalizados del documento"""
    
    # Estilo para títulos principales
    estilo_titulo = doc.styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
    estilo_titulo.font.name = 'Calibri'
    estilo_titulo.font.size = Pt(16)
    estilo_titulo.font.bold = True
    estilo_titulo.font.color.rgb = RGBColor(46, 84, 149)  # Azul oscuro
    estilo_titulo.paragraph_format.space_after = Pt(12)
    
    # Estilo para subtítulos
    estilo_subtitulo = doc.styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
    estilo_subtitulo.font.name = 'Calibri'
    estilo_subtitulo.font.size = Pt(14)
    estilo_subtitulo.font.bold = True
    estilo_subtitulo.font.color.rgb = RGBColor(64, 64, 64)  # Gris oscuro
    estilo_subtitulo.paragraph_format.space_before = Pt(18)
    estilo_subtitulo.paragraph_format.space_after = Pt(6)
    
    # Estilo para encabezados de sección
    estilo_seccion = doc.styles.add_style('Seccion', WD_STYLE_TYPE.PARAGRAPH)
    estilo_seccion.font.name = 'Calibri'
    estilo_seccion.font.size = Pt(12)
    estilo_seccion.font.bold = True
    estilo_seccion.paragraph_format.space_before = Pt(12)
    estilo_seccion.paragraph_format.space_after = Pt(6)
    
    # Estilo para código
    estilo_codigo = doc.styles.add_style('Codigo', WD_STYLE_TYPE.PARAGRAPH)
    estilo_codigo.font.name = 'Consolas'
    estilo_codigo.font.size = Pt(10)
    estilo_codigo.paragraph_format.left_indent = Inches(0.5)
    estilo_codigo.paragraph_format.space_before = Pt(6)
    estilo_codigo.paragraph_format.space_after = Pt(6)
    
    # Estilo normal modificado
    estilo_normal = doc.styles['Normal']
    estilo_normal.font.name = 'Calibri'
    estilo_normal.font.size = Pt(11)

def agregar_portada(doc):
    """Agrega la portada del documento"""
    
    # Título principal
    titulo = doc.add_heading('DOCUMENTO DE TRANSFERENCIA TECNOLÓGICA', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espacio
    doc.add_paragraph()
    
    # Información del sistema
    tabla_info = doc.add_table(rows=5, cols=2)
    tabla_info.style = 'LightShading-Accent1'
    
    datos_portada = [
        ("Sistema:", "[NOMBRE DEL SISTEMA WEB]"),
        ("Versión:", "[vX.X.X]"),
        ("Fecha de transferencia:", datetime.now().strftime('%d/%m/%Y')),
        ("Responsable actual:", "[Tu Nombre Completo]"),
        ("Área receptora:", "Operaciones e Infraestructura"),
    ]
    
    for i, (campo, valor) in enumerate(datos_portada):
        tabla_info.rows[i].cells[0].text = campo
        tabla_info.rows[i].cells[1].text = valor
        tabla_info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()

def agregar_indice(doc):
    """Agrega el índice del documento"""
    
    doc.add_heading('ÍNDICE', level=1)
    doc.add_paragraph()
    
    secciones = [
        "1. RESUMEN EJECUTIVO",
        "2. ARQUITECTURA DEL SISTEMA",
        "2.1 Diagrama de componentes",
        "2.2 Listado de microservicios",
        "3. INFRAESTRUCTURA DE PRODUCCIÓN",
        "4. PROCESO DE DESPLIEGUE ACTUAL",
        "5. CONFIGURACIÓN Y VARIABLES DE ENTORNO",
        "6. MONITOREO Y LOGS",
        "7. BACKUP Y RECUPERACIÓN",
        "8. SEGURIDAD",
        "9. PROCEDIMIENTOS DE INCIDENTES",
        "10. MEJORAS PLANEADAS / DEUDA TÉCNICA",
        "11. ANEXOS",
        "FIRMAS DE CONFORMIDAD"
    ]
    
    for seccion in secciones:
        p = doc.add_paragraph(seccion)
        p.paragraph_format.left_indent = Inches(0.5 if '.' in seccion[:3] else 0)
    
    doc.add_page_break()

def agregar_resumen_ejecutivo(doc):
    """Agrega la sección de resumen ejecutivo"""
    
    doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    
    # Crear tabla de resumen
    tabla = doc.add_table(rows=7, cols=2)
    tabla.style = 'LightGrid-Accent1'
    
    # Configurar ancho de columnas
    for row in tabla.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)
    
    datos_resumen = [
        ("Arquitectura:", "Sistema web basado en microservicios (25 servicios)"),
        ("Tecnología principal:", ".NET 8 (ASP.NET Core)"),
        ("Orquestación:", "Podman en producción"),
        ("Estrategia despliegue:", "CI/CD manual desde repositorio Git"),
        ("Entorno producción:", "[Especificar servidores/cloud]"),
        ("Disponibilidad:", "[Ej: 99.5% en horario comercial]"),
        ("SLA actual:", "[Especificar acuerdos de nivel de servicio]")
    ]
    
    for i, (campo, valor) in enumerate(datos_resumen):
        tabla.rows[i].cells[0].text = campo
        tabla.rows[i].cells[1].text = valor
        tabla.rows[i].cells[0].paragraphs[0].runs[0].bold = True

def agregar_arquitectura(doc):
    """Agrega la sección de arquitectura"""
    
    doc.add_heading('2. ARQUITECTURA DEL SISTEMA', level=1)
    
    # 2.1 Diagrama de componentes
    doc.add_heading('2.1 Diagrama de componentes', level=2)
    doc.add_paragraph('(Insertar diagrama arquitectónico aquí)', style='Intense Quote')
    
    # Diagrama ASCII
    p = doc.add_paragraph()
    run = p.add_run()
    run.text = """\
[API Gateway] → [Microservicios (25)] → [Bases de datos/Storage]
       ↑
[Cliente Web]  [Otros consumidores]\
"""
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    
    # 2.2 Listado de microservicios
    doc.add_heading('2.2 Listado de microservicios', level=2)
    
    # Crear tabla de microservicios
    tabla_ms = doc.add_table(rows=6, cols=5)
    tabla_ms.style = 'MediumShading1-Accent1'
    
    # Encabezados
    headers = ["#", "Nombre del Servicio", "Puerto", "Función principal", "Dependencias"]
    for i, header in enumerate(headers):
        cell = tabla_ms.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Datos de ejemplo (primeros 5 servicios)
    datos_ms = [
        ("1", "servicio-autenticacion", "5001", "Gestión de usuarios y tokens", "BD_Usuarios, Redis"),
        ("2", "servicio-catalogo", "5002", "Catálogo de productos", "BD_Catalogo"),
        ("3", "servicio-pedidos", "5003", "Procesamiento de pedidos", "BD_Pedidos, Redis"),
        ("4", "servicio-pagos", "5004", "Procesamiento de pagos", "API externa, BD_Pagos"),
        ("5", "servicio-notificaciones", "5005", "Envío de notificaciones", "SMTP, BD_Notificaciones")
    ]
    
    for i, datos in enumerate(datos_ms, 1):
        for j, valor in enumerate(datos):
            tabla_ms.rows[i].cells[j].text = str(valor)
    
    doc.add_paragraph("... (continuación para los 25 servicios)")

def agregar_infraestructura(doc):
    """Agrega la sección de infraestructura"""
    
    doc.add_heading('3. INFRAESTRUCTURA DE PRODUCCIÓN', level=1)
    
    # 3.1 Especificaciones del servidor
    doc.add_heading('3.1 Especificaciones del servidor', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run()
    run.text = """\
Servidor Principal:
  - Hostname: [nombre-servidor]
  - IP: [XXX.XXX.XXX.XXX]
  - SO: [Ej: RHEL 8.6 / Ubuntu 22.04]
  - CPU: [Especificaciones - ej: 8 cores]
  - RAM: [XX GB - ej: 32GB]
  - Storage: [XX GB - ej: 500GB] (Ruta principal: [/opt/aplicacion])
  - Podman version: [X.X.X - ej: 4.6.1]

Servidores Adicionales:
  - [Listar otros servidores si existen]\
"""
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

def agregar_despliegue(doc):
    """Agrega la sección de proceso de despliegue"""
    
    doc.add_heading('4. PROCESO DE DESPLIEGUE ACTUAL', level=1)
    
    # 4.1 Flujo completo
    doc.add_heading('4.1 Flujo completo', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run()
    run.text = """\
1. Desarrollo local → 2. Commit/Push a Git → 3. SSH al servidor → 
4. Git pull → 5. Build/Publicación → 6. Recrear contenedores → 
7. Health checks → 8. Validación\
"""
    run.font.name = 'Consolas'
    
    # 4.2 Comandos críticos
    doc.add_heading('4.2 Comandos críticos', level=2)
    
    comandos = [
        "# 1. Conectar al servidor",
        "ssh usuario@[IP-servidor] -p [puerto]",
        "",
        "# 2. Navegar al directorio",
        "cd /ruta/a/tu/proyecto",
        "",
        "# 3. Obtener últimos cambios",
        "git pull origin main",
        "",
        "# 4. Publicar servicio (ejemplo)",
        "dotnet publish servicio.csproj -c Release -o ./publish",
        "",
        "# 5. Recrear contenedor",
        "podman stop nombre-contenedor",
        "podman rm nombre-contenedor",
        "podman build -t imagen-tag .",
        "podman run -d -p 5001:5001 --name contenedor imagen-tag",
        "",
        "# 6. Verificar estado",
        "podman ps",
        "podman logs contenedor --tail 50",
        "",
        "# 7. Health check",
        "curl -f http://localhost:5001/health || echo 'FAILED'"
    ]
    
    for cmd in comandos:
        if cmd.startswith("#"):
            p = doc.add_paragraph(cmd)
            p.paragraph_format.left_indent = Inches(0)
        elif cmd:
            p = doc.add_paragraph(cmd, style='Codigo')
        else:
            doc.add_paragraph()

def agregar_configuracion(doc):
    """Agrega la sección de configuración"""
    
    doc.add_heading('5. CONFIGURACIÓN Y VARIABLES DE ENTORNO', level=1)
    
    # Tabla de variables
    tabla_vars = doc.add_table(rows=5, cols=3)  # 1 encabezado + 4 datos
    tabla_vars.style = 'LightGrid-Accent1'
    
    headers = ["Variable", "Ubicación", "Método de actualización"]
    for i, header in enumerate(headers):
        tabla_vars.rows[0].cells[i].text = header
        tabla_vars.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    datos_vars = [
        ("ConnectionStrings__Default", "HashiCorp Vault / Archivo encriptado", "Script de rotación mensual"),
        ("JWT__SecretKey", "Azure Key Vault", "Portal Azure + redeploy"),
        ("ExternalAPI__Key", "Variable de entorno en contenedor", "Update en podman run"),
        ("Logging__Level", "appsettings.Production.json", "Modificar y redeploy")
    ]
    
    for i, datos in enumerate(datos_vars, 1):
        for j, valor in enumerate(datos):
            tabla_vars.rows[i].cells[j].text = valor

def agregar_monitoreo(doc):
    """Agrega la sección de monitoreo"""
    
    doc.add_heading('6. MONITOREO Y LOGS', level=1)
    
    # 6.1 Métricas
    doc.add_heading('6.1 Métricas a monitorear', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run()
    run.text = """\
Críticas:
  - CPU uso > 80% por 5 min
  - Memoria uso > 85%
  - HTTP 5xx errors > 1%/min
  - Latencia p95 > 2s

Importantes:
  - Tasa de errores por servicio
  - Tiempo de respuesta promedio
  - Health checks fallidos
  - Espacio en disco < 20% libre\
"""
    run.font.name = 'Consolas'

def agregar_backup(doc):
    """Agrega la sección de backup"""
    
    doc.add_heading('7. BACKUP Y RECUPERACIÓN', level=1)
    
    # Tabla de estrategia de backup
    tabla_backup = doc.add_table(rows=5, cols=4)  # 1 encabezado + 4 datos
    tabla_backup.style = 'MediumList1-Accent1'
    
    headers = ["Componente", "Frecuencia", "Retención", "Ubicación"]
    for i, header in enumerate(headers):
        tabla_backup.rows[0].cells[i].text = header
        tabla_backup.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    datos_backup = [
        ("Bases de datos", "Diario (22:00)", "30 días", "/backups/db/"),
        ("Configuraciones", "Semanal (domingo)", "12 semanas", "/backups/config/"),
        ("Logs importantes", "Mensual", "1 año", "NAS/Cloud Storage"),
        ("Imágenes contenedores", "Por versión", "5 versiones", "Registry interno")
    ]
    
    for i, datos in enumerate(datos_backup, 1):
        for j, valor in enumerate(datos):
            tabla_backup.rows[i].cells[j].text = valor

def agregar_seguridad(doc):
    """Agrega la sección de seguridad"""
    
    doc.add_heading('8. SEGURIDAD', level=1)
    
    # Checklist de hardening
    doc.add_heading('8.1 Hardening aplicado', level=2)
    
    checklist = [
        ("✓", "Contenedores ejecutan como usuario no-root"),
        ("✓", "Secrets en variables de entorno (no en código)"),
        ("✓", "Firewall configurado (ufw/iptables)"),
        ("○", "Escaneo de vulnerabilidades periódico"),
        ("✓", "Logs de auditoría habilitados"),
        ("○", "WAF (Web Application Firewall) implementado")
    ]
    
    for estado, desc in checklist:
        p = doc.add_paragraph()
        p.add_run(f"{estado} {desc}").bold = (estado == "✓")

def agregar_incidentes(doc):
    """Agrega la sección de procedimientos de incidentes"""
    
    doc.add_heading('9. PROCEDIMIENTOS DE INCIDENTES', level=1)
    
    # Tabla de incidentes comunes
    tabla_incidentes = doc.add_table(rows=5, cols=4)  # 1 encabezado + 4 datos
    tabla_incidentes.style = 'LightShading-Accent1'
    
    headers = ["Síntoma", "Posible causa", "Acción inmediata", "Resolución"]
    for i, header in enumerate(headers):
        tabla_incidentes.rows[0].cells[i].text = header
        tabla_incidentes.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    datos_incidentes = [
        ("Error 502 en gateway", "Microservicio caído", "1. Verificar podman ps\n2. Revisar logs", "Restart del servicio"),
        ("Alta latencia", "CPU/Memoria saturada", "1. Usar top/htop\n2. Escalar temporalmente", "Optimizar o escalar recursos"),
        ("Conexión BD rechazada", "BD no responde", "1. Verificar proceso BD\n2. Check conexión", "Restart servicio BD"),
        ("Disk space full", "Logs sin rotación", "1. df -h\n2. Limpiar logs antiguos", "Implementar log rotation")
    ]
    
    for i, datos in enumerate(datos_incidentes, 1):
        for j, valor in enumerate(datos):
            tabla_incidentes.rows[i].cells[j].text = str(valor)

def agregar_mejoras(doc):
    """Agrega la sección de mejoras planeadas"""
    
    doc.add_heading('10. MEJORAS PLANEADAS / DEUDA TÉCNICA', level=1)
    
    tabla_mejoras = doc.add_table(rows=6, cols=4)  # 1 encabezado + 5 datos
    tabla_mejoras.style = 'MediumGrid3-Accent1'
    
    headers = ["Item", "Prioridad", "Estimado", "Notas"]
    for i, header in enumerate(headers):
        tabla_mejoras.rows[0].cells[i].text = header
        tabla_mejoras.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    datos_mejoras = [
        ("Implementar CI/CD automático", "Alta", "2-3 sprints", "Jenkins/GitHub Actions"),
        ("Migrar a Kubernetes", "Media", "Q3 2024", "Evaluar costos/beneficios"),
        ("Centralizar logs con ELK", "Alta", "1 sprint", "Mejora debugging"),
        ("Autoscaling horizontal", "Baja", "Q4 2024", "Depende de crecimiento tráfico"),
        ("Monitoring avanzado", "Media", "Q2 2024", "Prometheus + Grafana dashboards")
    ]
    
    for i, datos in enumerate(datos_mejoras, 1):
        for j, valor in enumerate(datos):
            tabla_mejoras.rows[i].cells[j].text = valor

def agregar_anexos(doc):
    """Agrega la sección de anexos"""
    
    doc.add_heading('11. ANEXOS', level=1)
    
    # Checklist pre-despliegue
    doc.add_heading('A. Checklist pre-despliegue', level=2)
    
    checklist_items = [
        "Backups completados",
        "Team notificado",
        "Ventana de mantenimiento confirmada",
        "Rollback plan listo",
        "Health checks configurados",
        "Documentación actualizada",
        "Tests de integración pasados"
    ]
    
    for item in checklist_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"[ ] {item}")

def agregar_firmas(doc):
    """Agrega la sección de firmas"""
    
    doc.add_page_break()
    doc.add_heading('FIRMAS DE CONFORMIDAD', level=1)
    
    # Tabla de firmas
    tabla_firmas = doc.add_table(rows=5, cols=4)  # 1 encabezado + 4 datos
    tabla_firmas.style = 'LightGrid-Accent1'
    
    headers = ["Rol", "Nombre", "Firma", "Fecha"]
    for i, header in enumerate(headers):
        tabla_firmas.rows[0].cells[i].text = header
        tabla_firmas.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    datos_firmas = [
        ("Entrega (Desarrollo)", "[Tu Nombre Completo]", "__________", datetime.now().strftime('%d/%m/%Y')),
        ("Recibe (Operaciones)", "[Nombre Receptor]", "__________", ""),
        ("Testigo (Infraestructura)", "[Nombre Testigo]", "__________", ""),
        ("Aprobación (Gerencia)", "[Nombre Gerente]", "__________", "")
    ]
    
    for i, datos in enumerate(datos_firmas, 1):
        for j, valor in enumerate(datos):
            tabla_firmas.rows[i].cells[j].text = str(valor)
    
    # Notas finales
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("NOTA: ").bold = True
    p.add_run("Este documento debe actualizarse con cada cambio arquitectónico significativo.")
    
    p = doc.add_paragraph()
    p.add_run("Última revisión: ").bold = True
    p.add_run(datetime.now().strftime('%d/%m/%Y'))
    
    p = doc.add_paragraph()
    p.add_run("Próxima revisión programada: ").bold = True
    p.add_run("[DD/MM/AAAA]")
    
    p = doc.add_paragraph()
    p.add_run("Custodio del documento: ").bold = True
    p.add_run("[Nombre del Arquitecto/Lead Técnico]")

# Ejecutar el generador
if __name__ == "__main__":
    print("🚀 Generando documento de transferencia tecnológica...")
    archivo_generado = crear_documento_transferencia()
    print("🎉 Documento generado exitosamente!")
    print("\n📋 Pasos siguientes:")
    print("1. Revisar el documento generado")
    print("2. Completar los campos entre [corchetes]")
    print("3. Insertar diagramas en las secciones indicadas")
    print("4. Personalizar según necesidades específicas")