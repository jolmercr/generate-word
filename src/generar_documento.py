#!/usr/bin/env python3
"""
Script de ejemplo para generar documentos de Word.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def crear_documento_ejemplo():
    """Crea un documento de Word de ejemplo."""
    # Crear un nuevo documento
    doc = Document()
    
    # Agregar un título
    titulo = doc.add_heading('Documento de Ejemplo', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar un subtítulo
    doc.add_heading('Sección 1: Introducción', level=1)
    
    # Agregar párrafos
    doc.add_paragraph(
        'Este es un documento de ejemplo creado con python-docx. '
        'Puedes personalizar el contenido según tus necesidades.'
    )
    
    # Agregar una lista con viñetas
    doc.add_paragraph('Características principales:', style='Heading 2')
    doc.add_paragraph('Crear documentos de Word programáticamente', style='List Bullet')
    doc.add_paragraph('Agregar texto, tablas e imágenes', style='List Bullet')
    doc.add_paragraph('Personalizar estilos y formato', style='List Bullet')
    
    # Agregar una tabla
    doc.add_heading('Sección 2: Tabla de Ejemplo', level=1)
    tabla = doc.add_table(rows=3, cols=3)
    tabla.style = 'Light Grid Accent 1'
    
    # Encabezados de la tabla
    celdas_encabezado = tabla.rows[0].cells
    celdas_encabezado[0].text = 'Producto'
    celdas_encabezado[1].text = 'Cantidad'
    celdas_encabezado[2].text = 'Precio'
    
    # Datos de ejemplo
    datos = [
        ('Producto A', '10', '$100'),
        ('Producto B', '5', '$50')
    ]
    
    for i, (producto, cantidad, precio) in enumerate(datos, start=1):
        celdas = tabla.rows[i].cells
        celdas[0].text = producto
        celdas[1].text = cantidad
        celdas[2].text = precio
    
    # Agregar un salto de página
    doc.add_page_break()
    
    # Nueva sección
    doc.add_heading('Sección 3: Conclusión', level=1)
    doc.add_paragraph(
        'Este es un ejemplo básico de las capacidades de python-docx. '
        'Puedes expandir este código para crear documentos más complejos.'
    )
    
    # Guardar el documento
    nombre_archivo = 'documento_ejemplo.docx'
    doc.save(nombre_archivo)
    print(f'✓ Documento creado exitosamente: {nombre_archivo}')


if __name__ == '__main__':
    crear_documento_ejemplo()
